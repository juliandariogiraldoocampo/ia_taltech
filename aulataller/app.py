import streamlit as st
import pandas as pd
from groq import Groq
from docx import Document
from io import BytesIO
from datetime import datetime

st.set_page_config(page_title="Informe IA", page_icon="📊")
st.title("📊 Excel → Informe Word")

# == API KEY ==============================================
# En .streamlit/secrets.toml: GROQ_API_KEY = "tu-clave"
cliente = Groq(api_key=st.secrets['GROQ_API_KEY'])

archivo = st.file_uploader("Sube tu Excel", type=['xlsx', 'xls'])
contexto = st.text_area("Contexto (opcional)", height=100)

modelo = st.selectbox("Modelo", ["openai/gpt-oss-120b", "llama-3.3-70b-versatile"])

if archivo and st.button("Generar informe"):
    with st.spinner("Procesando..."):
        datos = pd.read_excel(archivo, sheet_name=None)

        resumen = ""
        for nombre, hoja in datos.items():
            resumen += f"\nHoja: {nombre}\n"
            resumen += f"Filas: {len(hoja)}\n"
            resumen += f"Columnas: {', '.join(hoja.columns[:5])}\n"

        prompt = f"""
        Contexto: {contexto if contexto else 'Datos para análisis'}

        Datos del Excel:
        {resumen}
        Utiliza SOLO los datos del archivo Excel para generar el informe. El contexto de los datos son aeropuertos en Colombia.
        Genera un informe con:
        1. Resumen ejecutivo (400 palabras)
        2. 3 insights clave (con ejemplos)
        3. 2 recomendaciones (con acciones específicas)

        IMPORTANTE: No uses Markdown. Para títulos, negritas ni cursivas. Usa solo texto plano con saltos de línea.
        """

        respuesta = cliente.chat.completions.create(
            #model='llama-3.3-70b-versatile',
            model=modelo,
            messages=[{'role': 'user', 'content': prompt}],
            temperature=0.2,
        )

        texto = respuesta.choices[0].message.content or "No se pudo generar el informe"

        # Crear Word
        doc = Document()
        doc.add_heading('Informe Automático', 0)
        doc.add_paragraph(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        doc.add_paragraph()
        doc.add_paragraph(texto)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Informe listo")
        st.download_button(
            "📥 Descargar Word",
            buffer,
            file_name=f"informe_{datetime.now().strftime('%Y%m%d')}.docx"
        )

        with st.expander("Ver informe"):
            st.write(texto)